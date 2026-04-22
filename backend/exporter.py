import os
from fpdf import FPDF
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def export_to_word(handbook_text: str, topic: str) -> str:
    doc = Document()

    title = doc.add_heading(f"Comprehensive Handbook: {topic}", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    lines = handbook_text.split("\n")
    for line in lines:
        line = line.strip()
        if not line:
            continue
        if line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("---"):
            doc.add_paragraph("─" * 50)
        else:
            para = doc.add_paragraph(line)
            para.style.font.size = Pt(11)

    output_path = f"outputs/handbook_{topic[:30].replace(' ', '_')}.docx"
    os.makedirs("outputs", exist_ok=True)
    doc.save(output_path)
    print(f"Word document saved to {output_path}")
    return output_path


def export_to_pdf(handbook_text: str, topic: str) -> str:
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_font("Helvetica", size=12)

    lines = handbook_text.split("\n")
    for line in lines:
        line = line.strip()
        if not line:
            pdf.ln(4)
            continue

        clean_line = line.encode("latin-1", errors="replace").decode("latin-1")

        if line.startswith("# "):
            pdf.set_font("Helvetica", "B", 18)
            try:
                pdf.multi_cell(0, 10, clean_line[2:])
            except Exception:
                pass
            pdf.ln(4)
            pdf.set_font("Helvetica", size=12)
        elif line.startswith("## "):
            pdf.set_font("Helvetica", "B", 15)
            try:
                pdf.multi_cell(0, 8, clean_line[3:])
            except Exception:
                pass
            pdf.ln(3)
            pdf.set_font("Helvetica", size=12)
        elif line.startswith("### "):
            pdf.set_font("Helvetica", "B", 13)
            try:
                pdf.multi_cell(0, 7, clean_line[4:])
            except Exception:
                pass
            pdf.ln(2)
            pdf.set_font("Helvetica", size=12)
        elif line.startswith("---"):
            pdf.ln(2)
        else:
            try:
                pdf.multi_cell(0, 7, clean_line)
            except Exception:
                pass

    output_path = f"outputs/handbook_{topic[:30].replace(' ', '_')}.pdf"
    os.makedirs("outputs", exist_ok=True)
    pdf.output(output_path)
    print(f"PDF saved to {output_path}")
    return output_path